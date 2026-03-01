"""
Foreperson.ai – Construction Document Parser

Supported formats: PDF, DOCX, XLSX/XLS, CSV, TXT, PNG, JPG, JPEG, GIF, TIFF, WEBP, BMP
Image extraction priority: Vision API (GPT-4o / Claude) → Tesseract OCR → placeholder
"""

from __future__ import annotations

import base64
import io
import logging
import os
import re
from pathlib import Path
from typing import Dict, List, Optional

logger = logging.getLogger(__name__)

# ── Optional library detection ────────────────────────────────

try:
    import fitz        # pymupdf
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False

try:
    import pdfplumber
    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    import PyPDF2
    HAS_PYPDF2 = True
except ImportError:
    HAS_PYPDF2 = False

try:
    import pytesseract
    from PIL import Image, ImageEnhance, ImageFilter
    HAS_TESSERACT = True
except ImportError:
    HAS_TESSERACT = False

try:
    from PIL import Image as PILImage
    HAS_PILLOW = True
except ImportError:
    HAS_PILLOW = False

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import openpyxl
    HAS_OPENPYXL = True
except ImportError:
    HAS_OPENPYXL = False

try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

# ── Document type registry ────────────────────────────────────
# Keys must match the frontend CATS registry (lowercase snake_case)

DOCUMENT_TYPES: Dict[str, Dict] = {
    'rfi': {
        'keywords': ['rfi', 'request for information', 'clarification requested', 'field question'],
        'patterns':  [r'RFI[-\s#]?\d+', r'REQUEST FOR INFORMATION'],
        'filename':  ['rfi'],
    },
    'submittal': {
        'keywords': ['submittal', 'shop drawing', 'product data', 'samples', 'mock-up', 'submittal log'],
        'patterns':  [r'SUBMITTAL[-\s]?\d+', r'SHOP DRAWING', r'PRODUCT DATA'],
        'filename':  ['submittal', 'shop_draw', 'shopdraw'],
    },
    'specification': {
        'keywords': ['specification', 'division', 'csi', 'masterformat', 'part 1 general',
                     'part 2 products', 'part 3 execution'],
        'patterns':  [r'SECTION \d{5}', r'DIVISION \d+', r'\b\d{2} \d{2} \d{2}\b'],
        'filename':  ['spec', 'div0', 'div1', 'div2', 'div3', 'div4', 'div5',
                      'masterformat', 'csi'],
    },
    'contract': {
        'keywords': ['contract', 'agreement', 'general conditions', 'subcontract',
                     'whereas', 'witnesseth', 'scope of work', 'liquidated damages'],
        'patterns':  [r'CONTRACT DOCUMENTS', r'GENERAL CONDITIONS', r'SUBCONTRACT AGREEMENT'],
        'filename':  ['contract', 'subcontract', 'agreement', 'gcmax', 'aia_a'],
    },
    'change_order': {
        'keywords': ['change order', 'pco', 'potential change order', 'change directive',
                     'field order', 'bulletin', 'asi', 'supplemental instruction'],
        'patterns':  [r'CO[-\s#]?\d+', r'PCO[-\s#]?\d+', r'CHANGE ORDER \d+',
                      r'CHANGE DIRECTIVE', r'ASI[-\s]?\d+'],
        'filename':  ['change_order', 'change order', 'co_', 'pco_', 'asi_', 'bulletin'],
    },
    'floor_plan': {
        'keywords': ['floor plan', 'reflected ceiling plan', 'rcp', 'room schedule',
                     'partition', 'interior layout', 'furniture layout'],
        'patterns':  [r'FLOOR PLAN', r'LEVEL \d+', r'REFLECTED CEILING', r'RCP',
                      r'[A-Z]\d?\.\d{3}'],
        'filename':  ['floor', 'flr', 'fp_', 'rcp', 'ceiling_plan'],
    },
    'site_plan': {
        'keywords': ['site plan', 'civil', 'grading plan', 'drainage', 'utilities',
                     'topographic', 'survey', 'erosion control', 'stormwater'],
        'patterns':  [r'SITE PLAN', r'GRADING PLAN', r'C[-\s]?\d+', r'CIVIL SHEET'],
        'filename':  ['site', 'civil', 'grading', 'survey', 'topog', 'drainage'],
    },
    'drawing': {
        'keywords': ['drawing', 'elevation', 'section', 'detail', 'scale', 'north arrow',
                     'structural', 'mechanical', 'electrical', 'plumbing', 'fire protection'],
        'patterns':  [r'[A-Z]-\d+', r'SHEET \d+ OF \d+', r'DWG[-\s]?\d+',
                      r'DRAWING NO\.?', r'REV\.?\s*\d+'],
        'filename':  ['dwg', 'drawing', 'elevation', 'section', 'detail', 'struct',
                      'mech', 'elec', 'plumb'],
    },
    'schedule': {
        'keywords': ['schedule', 'timeline', 'milestone', 'gantt', 'baseline',
                     'lookahead', 'critical path', 'activity id', 'duration', 'float'],
        'patterns':  [r'MASTER SCHEDULE', r'BASELINE SCHEDULE', r'\d+-DAY LOOKAHEAD',
                      r'CRITICAL PATH'],
        'filename':  ['schedule', 'sched', 'timeline', 'lookahead', 'gantt', 'p6'],
    },
    'budget': {
        'keywords': ['budget', 'cost report', 'pay application', 'schedule of values',
                     'sov', 'aia g702', 'aia g703', 'earned value', 'cost-to-complete'],
        'patterns':  [r'PAY APPLICATION', r'G702', r'G703', r'SCHEDULE OF VALUES',
                      r'COST REPORT'],
        'filename':  ['budget', 'cost_report', 'pay_app', 'payapp', 'sov', 'g702'],
    },
    'report': {
        'keywords': ['daily report', 'inspection report', 'safety report', 'progress report',
                     'site observation', 'testing report', 'soils report', 'rca'],
        'patterns':  [r'DAILY REPORT', r'INSPECTION REPORT', r'SAFETY REPORT',
                      r'PROGRESS REPORT'],
        'filename':  ['report', 'daily', 'inspection', 'safety', 'progress', 'testing'],
    },
    'minutes': {
        'keywords': ['meeting minutes', 'minutes of meeting', 'action items', 'attendees',
                     'agenda', 'mom', 'next meeting', 'decisions made'],
        'patterns':  [r'MEETING MINUTES', r'MINUTES OF', r'ACTION ITEMS'],
        'filename':  ['minutes', 'meeting', 'mom_', 'oac_'],
    },
    'punch_list': {
        'keywords': ['punch list', 'deficiency list', 'punchlist', 'closeout',
                     'substantial completion', 'certificate of occupancy', 'co inspection'],
        'patterns':  [r'PUNCH LIST', r'DEFICIENCY LIST', r'CLOSEOUT', r'SUBSTANTIAL COMPLETION'],
        'filename':  ['punch', 'punchlist', 'deficiency', 'closeout', 'substantial'],
    },
}

IMAGE_EXTENSIONS = {'.png', '.jpg', '.jpeg', '.gif', '.tiff', '.tif', '.webp', '.bmp'}
VISION_PROMPT = """You are analyzing a construction document image. Extract ALL information present including:

- Every piece of text visible (titles, labels, notes, dimensions, callouts, room names, material specs)
- Drawing numbers, sheet numbers, revision numbers
- Scale information and north arrow direction
- Title block contents (project name, date, drawn by, checked by, etc.)
- All dimension strings and measurements
- Door/window tags and schedule references
- Material and finish notations
- Section marks, detail bubbles, elevation markers
- Grid lines and column labels
- Any tables or schedules visible
- Keynotes or legend items

Format the output as structured text that preserves all information. If this is a floor plan, site plan, or structural drawing, describe the layout and key elements in addition to extracting text.

Do not interpret or explain — just extract and organize all visible information."""


# ── Vision API image extraction ───────────────────────────────

def _image_to_base64(file_path: str) -> tuple[str, str]:
    """Return (base64_data, media_type) for an image file."""
    ext = Path(file_path).suffix.lower()
    mime = {
        '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
        '.png': 'image/png', '.gif': 'image/gif',
        '.webp': 'image/webp', '.tiff': 'image/tiff',
        '.tif': 'image/tiff', '.bmp': 'image/bmp',
    }.get(ext, 'image/jpeg')
    with open(file_path, 'rb') as f:
        return base64.b64encode(f.read()).decode('utf-8'), mime


def _vision_openai(file_path: str) -> Optional[str]:
    """Extract text from image using OpenAI vision (gpt-4o)."""
    api_key = os.environ.get('OPENAI_API_KEY')
    if not api_key:
        return None
    try:
        from openai import OpenAI
        b64, mime = _image_to_base64(file_path)
        client = OpenAI(api_key=api_key)
        resp = client.chat.completions.create(
            model='gpt-4o',
            max_tokens=2000,
            messages=[{
                'role': 'user',
                'content': [
                    {'type': 'text', 'text': VISION_PROMPT},
                    {'type': 'image_url', 'image_url': {'url': f'data:{mime};base64,{b64}', 'detail': 'high'}},
                ],
            }],
        )
        return resp.choices[0].message.content
    except Exception as e:
        logger.warning('OpenAI vision failed: %s', e)
        return None


def _vision_anthropic(file_path: str) -> Optional[str]:
    """Extract text from image using Anthropic vision (claude-sonnet-4-6)."""
    api_key = os.environ.get('ANTHROPIC_API_KEY')
    if not api_key:
        return None
    try:
        import anthropic
        b64, mime = _image_to_base64(file_path)
        client = anthropic.Anthropic(api_key=api_key)
        resp = client.messages.create(
            model='claude-sonnet-4-6',
            max_tokens=2000,
            messages=[{
                'role': 'user',
                'content': [
                    {'type': 'image', 'source': {'type': 'base64', 'media_type': mime, 'data': b64}},
                    {'type': 'text', 'text': VISION_PROMPT},
                ],
            }],
        )
        return resp.content[0].text
    except Exception as e:
        logger.warning('Anthropic vision failed: %s', e)
        return None


def _vision_tesseract(file_path: str) -> Optional[str]:
    """Fallback: extract text from image using Tesseract OCR."""
    if not HAS_TESSERACT:
        return None
    try:
        img = PILImage.open(file_path).convert('L')
        img = ImageEnhance.Contrast(img).enhance(1.5)
        img = ImageEnhance.Sharpness(img).enhance(1.3)
        img = img.filter(ImageFilter.MedianFilter(size=3))
        text = pytesseract.image_to_string(img, config='--oem 3 --psm 6')
        if not text.strip():
            text = pytesseract.image_to_string(img, config='--oem 3 --psm 11')
        return text.strip() or None
    except Exception as e:
        logger.warning('Tesseract failed: %s', e)
        return None


def extract_text_from_image(file_path: str) -> str:
    """Extract text/content from an image file using best available method."""
    logger.info('Extracting image: %s', Path(file_path).name)

    # Vision API is far superior for construction drawings
    text = _vision_openai(file_path) or _vision_anthropic(file_path)
    if text:
        logger.info('Image extracted via vision API (%d chars)', len(text))
        return text

    # Tesseract fallback
    text = _vision_tesseract(file_path)
    if text:
        logger.info('Image extracted via Tesseract (%d chars)', len(text))
        return text

    return '[Image uploaded — no text extraction available. Install an API key for vision analysis.]'


# ── Type detection ────────────────────────────────────────────

def detect_document_type(text: str, filename: str = '') -> str:
    """Score each document type and return the best match (lowercase snake_case)."""
    text_upper = text.upper()
    filename_upper = filename.upper()
    scores: Dict[str, int] = {}

    for doc_type, rules in DOCUMENT_TYPES.items():
        score = 0
        for kw in rules.get('keywords', []):
            if kw.upper() in text_upper:
                score += 2
        for pattern in rules.get('patterns', []):
            if re.search(pattern, text_upper):
                score += 3
        for hint in rules.get('filename', []):
            if hint.upper() in filename_upper:
                score += 4  # filename is a strong signal
        scores[doc_type] = score

    if not scores:
        return 'unknown'
    best_type, best_score = max(scores.items(), key=lambda x: x[1])
    return best_type if best_score > 0 else 'unknown'


# ── Per-format extractors ─────────────────────────────────────

def _extract_pdf(file_path: str) -> str:
    if HAS_PYMUPDF:
        try:
            doc = fitz.open(file_path)
            pages_text: List[str] = []
            ocr_needed: List[tuple] = []

            for i, page in enumerate(doc, 1):
                txt = page.get_text()
                if txt.strip() and len(txt.strip()) >= 50:
                    pages_text.append(f'--- Page {i} ---\n{txt}')
                else:
                    pages_text.append(f'--- Page {i} ---\n[scanning...]')
                    ocr_needed.append((i, page))

            # OCR scanned pages
            if ocr_needed:
                for page_num, page in ocr_needed:
                    mat = fitz.Matrix(400 / 72, 400 / 72)
                    pix = page.get_pixmap(matrix=mat, alpha=False)
                    img_data = pix.tobytes('png')
                    if HAS_TESSERACT:
                        img = PILImage.open(io.BytesIO(img_data)).convert('L')
                        img = ImageEnhance.Contrast(img).enhance(1.5)
                        ocr_text = pytesseract.image_to_string(img, config='--oem 3 --psm 6')
                        if not ocr_text.strip():
                            ocr_text = pytesseract.image_to_string(img, config='--oem 3 --psm 11')
                        replacement = f'--- Page {page_num} (OCR) ---\n{ocr_text.strip()}'
                    else:
                        replacement = f'--- Page {page_num} ---\n[Scanned page — install tesseract for OCR]'
                    pages_text[page_num - 1] = replacement

            doc.close()
            result = '\n'.join(pages_text)
            return result if result.strip() else 'No text found in PDF'
        except Exception as e:
            logger.warning('pymupdf error: %s', e)

    if HAS_PDFPLUMBER:
        try:
            text = ''
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    txt = page.extract_text() or ''
                    text += f'\n--- Page {i} ---\n{txt}'
                    tables = page.extract_tables()
                    for table in tables:
                        for row in table:
                            if row:
                                text += ' | '.join(str(c) if c else '' for c in row) + '\n'
            return text or 'No text found in PDF'
        except Exception as e:
            logger.warning('pdfplumber error: %s', e)

    if HAS_PYPDF2:
        try:
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                return '\n'.join(
                    f'--- Page {i+1} ---\n{p.extract_text() or "[no text]"}'
                    for i, p in enumerate(reader.pages)
                )
        except Exception as e:
            logger.error('PyPDF2 error: %s', e)

    return 'Error: No PDF extraction library available'


def _extract_docx(file_path: str) -> str:
    import tempfile
    import zipfile

    parts: List[str] = []

    # 1. Extract text from paragraphs and tables
    if HAS_DOCX:
        try:
            doc = DocxDocument(file_path)
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = '\t'.join(c.text for c in row.cells if c.text.strip())
                    if row_text.strip():
                        parts.append(row_text)
        except Exception as e:
            logger.warning('DOCX text extraction error: %s', e)

    # 2. Extract embedded images (DOCX is a ZIP — images live in word/media/)
    try:
        with zipfile.ZipFile(file_path, 'r') as z:
            media = [n for n in z.namelist() if n.startswith('word/media/')]
            if media:
                logger.info('DOCX contains %d embedded image(s)', len(media))
            for img_name in media:
                ext = Path(img_name).suffix.lower()
                # Skip vector/meta formats Tesseract/vision can't handle
                if ext in ('.wmf', '.emf'):
                    parts.append(f'[Embedded vector graphic: {Path(img_name).name} — cannot extract text]')
                    continue
                img_data = z.read(img_name)
                with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as tmp:
                    tmp.write(img_data)
                    tmp_path = tmp.name
                try:
                    img_text = extract_text_from_image(tmp_path)
                    if img_text:
                        parts.append(f'[Embedded image: {Path(img_name).name}]\n{img_text}')
                finally:
                    try:
                        os.unlink(tmp_path)
                    except OSError:
                        pass
    except zipfile.BadZipFile:
        pass  # Not a valid zip/docx
    except Exception as e:
        logger.warning('DOCX image extraction error: %s', e)

    if not parts:
        return '[Document contains no extractable text. It may consist entirely of embedded graphics that could not be read.]'
    return '\n'.join(parts)


def _extract_excel(file_path: str) -> str:
    if HAS_PANDAS:
        try:
            xl = pd.read_excel(file_path, sheet_name=None)
            parts: List[str] = []
            for sheet_name, df in xl.items():
                parts.append(f'Sheet: {sheet_name}')
                parts.append(df.to_string(index=False))
                parts.append('')
            return '\n'.join(parts)
        except Exception as e:
            logger.warning('pandas excel error: %s', e)

    if HAS_OPENPYXL:
        try:
            wb = openpyxl.load_workbook(file_path, data_only=True)
            parts: List[str] = []
            for name in wb.sheetnames:
                ws = wb[name]
                parts.append(f'Sheet: {name}')
                for row in ws.iter_rows(values_only=True):
                    parts.append('\t'.join(str(c) if c is not None else '' for c in row))
            return '\n'.join(parts)
        except Exception as e:
            logger.error('openpyxl error: %s', e)

    return 'Error: No Excel library available'


def _extract_csv(file_path: str) -> str:
    if HAS_PANDAS:
        try:
            df = pd.read_csv(file_path, encoding='utf-8', errors='replace')
            return df.to_string(index=False)
        except Exception as e:
            logger.warning('pandas csv error: %s', e)
    try:
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    except Exception as e:
        return f'Error reading CSV: {e}'


# ── Main parser class ─────────────────────────────────────────

class ConstructionDocumentParser:
    """Parse construction documents of any supported format."""

    SUPPORTED = {
        '.pdf', '.docx', '.doc',
        '.xlsx', '.xls',
        '.csv',
        '.txt',
        *IMAGE_EXTENSIONS,
    }

    def parse_document(self, file_path: str) -> Dict:
        fp = Path(file_path)
        if not fp.exists():
            return {'error': 'File not found'}

        ext = fp.suffix.lower()
        if ext not in self.SUPPORTED:
            return {'error': f'Unsupported file type: {ext}'}

        if ext == '.pdf':
            text = _extract_pdf(str(fp))
        elif ext in ('.docx', '.doc'):
            text = _extract_docx(str(fp))
        elif ext in ('.xlsx', '.xls'):
            text = _extract_excel(str(fp))
        elif ext == '.csv':
            text = _extract_csv(str(fp))
        elif ext == '.txt':
            try:
                text = fp.read_text(encoding='utf-8', errors='replace')
            except Exception as e:
                return {'error': f'Could not read text file: {e}'}
        elif ext in IMAGE_EXTENSIONS:
            text = extract_text_from_image(str(fp))
        else:
            return {'error': f'Unsupported file type: {ext}'}

        doc_type = detect_document_type(text, fp.name)

        return {
            'filename': fp.name,
            'file_path': str(fp),
            'document_type': doc_type,
            'text_content': text,
            'word_count': len(text.split()),
            'char_count': len(text),
        }

    def parse_directory(self, directory_path: str) -> List[Dict]:
        directory = Path(directory_path)
        docs: List[Dict] = []
        for fp in directory.rglob('*'):
            if fp.is_file() and fp.suffix.lower() in self.SUPPORTED:
                logger.info('Parsing: %s', fp.name)
                docs.append(self.parse_document(str(fp)))
        return docs

    # Legacy alias used by some callers
    def identify_document_type(self, text: str, filename: str = '') -> str:
        return detect_document_type(text, filename)
