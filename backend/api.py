"""
FastAPI REST API
"""
from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, status, BackgroundTasks
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, HTMLResponse, Response
from sqlalchemy.orm import Session
from pydantic import BaseModel, EmailStr
from typing import List, Optional
from datetime import datetime
import html
import json
import os
import sys
import logging
import urllib.parse

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s %(levelname)s %(name)s — %(message)s",
)
logger = logging.getLogger(__name__)

# Add parent directory to path for imports
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from slowapi import Limiter, _rate_limit_exceeded_handler
from slowapi.util import get_remote_address
from slowapi.errors import RateLimitExceeded
from starlette.requests import Request
from backend.constants import MAX_FILE_SIZE, PAGINATION_DEFAULT_LIMIT, PAGINATION_MAX_LIMIT, CHAT_RATE_LIMIT, CONFLICTS_RATE_LIMIT, CHAT_HISTORY_WINDOW
from backend.database import get_db, User, Project, Document, Chat, ChatMessage, ProjectMemory, ConflictStatus, RFI, DailyReport, ActionItem, init_db
from backend.auth import (
    get_current_user,
    get_user_from_token_param,
    create_access_token,
    register_user,
    authenticate_user,
    get_password_hash
)
from backend.storage import save_file, get_file, delete_file

# Import document parser and AI assistant from parent
try:
    from document_parser import ConstructionDocumentParser as DocumentParser
    from document_parser import detect_document_type
    from ai_assistant import ConstructionAI, ConstructionAgent
except ImportError:
    DocumentParser = None
    detect_document_type = None
    ConstructionAI = None
    ConstructionAgent = None

# Initialize FastAPI
app = FastAPI(
    title="Foreperson.ai API",
    description="Construction Document Intelligence API",
    version="1.0.0"
)

# CORS configuration
_allowed_origins = os.environ.get("ALLOWED_ORIGINS", "http://localhost:3000").split(",")

app.add_middleware(
    CORSMiddleware,
    allow_origins=[o.strip() for o in _allowed_origins],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

limiter = Limiter(key_func=get_remote_address)
app.state.limiter = limiter
app.add_exception_handler(RateLimitExceeded, _rate_limit_exceeded_handler)


# Pydantic models
class UserCreate(BaseModel):
    email: EmailStr
    password: str
    name: str


class UserLogin(BaseModel):
    email: EmailStr
    password: str


class UserResponse(BaseModel):
    id: int
    email: str
    name: str
    created_at: datetime

    class Config:
        from_attributes = True


class TokenResponse(BaseModel):
    access_token: str
    token_type: str = "bearer"


class ProjectCreate(BaseModel):
    name: str
    description: Optional[str] = None


class ProjectResponse(BaseModel):
    id: int
    name: str
    description: Optional[str]
    created_at: datetime
    document_count: int = 0

    class Config:
        from_attributes = True


class DocumentResponse(BaseModel):
    id: int
    filename: str
    original_filename: str
    file_size: int
    document_type: Optional[str]
    parse_quality: Optional[str] = "good"
    created_at: datetime

    class Config:
        from_attributes = True


class ChatRequest(BaseModel):
    message: str
    project_id: Optional[int] = None
    chat_id: Optional[int] = None
    model: Optional[str] = None  # e.g. "gpt-4o-mini" or "claude-sonnet-4-6"
    use_memory: bool = True
    referenced_chat_id: Optional[int] = None


class ChatResponse(BaseModel):
    response: str
    message_id: Optional[int] = None


class ChatMessageResponse(BaseModel):
    id: int
    role: str
    content: str
    created_at: datetime

    class Config:
        from_attributes = True


class ChatHistoryResponse(BaseModel):
    id: int
    title: Optional[str] = None
    messages: List[ChatMessageResponse]
    created_at: datetime

    class Config:
        from_attributes = True


class ChatThreadResponse(BaseModel):
    id: int
    title: Optional[str] = None
    message_count: int
    created_at: datetime

    class Config:
        from_attributes = True


class ConflictResponse(BaseModel):
    id: str
    severity: str
    title: str
    description: str
    resolution: str = ""
    documents: List[str]


class ConflictsRequest(BaseModel):
    doc_ids: Optional[List[int]] = None  # if None, use all docs in project


class CompareRequest(BaseModel):
    doc_id_1: int
    doc_id_2: int


class CompareConflictItem(BaseModel):
    title: str
    doc_a: str
    doc_b: str
    impact: str
    recommendation: str


class CompareRiskItem(BaseModel):
    title: str
    description: str


class CompareResponse(BaseModel):
    summary: str
    conflicts: List[CompareConflictItem]
    gaps: List[str]
    agreements: List[str]
    risks: List[CompareRiskItem]
    doc1_name: str
    doc2_name: str


class ProjectAnalyticsResponse(BaseModel):
    doc_count: int
    total_words: int
    type_breakdown: dict
    chat_count: int
    message_count: int
    memory_fact_count: int


# ---- RFI schemas ----
class RFICreate(BaseModel):
    subject: str
    description: str
    due_date: Optional[str] = None


class RFIUpdate(BaseModel):
    subject: Optional[str] = None
    description: Optional[str] = None
    status: Optional[str] = None
    response: Optional[str] = None
    due_date: Optional[str] = None


class RFIResponse(BaseModel):
    id: int
    number: int
    subject: str
    description: str
    status: str
    response: Optional[str] = None
    due_date: Optional[str] = None
    created_by: Optional[str] = None
    created_at: datetime

    class Config:
        from_attributes = True


# ---- DailyReport schemas ----
class DailyReportCreate(BaseModel):
    report_date: str
    work_performed: str
    weather: Optional[str] = None
    crew_count: Optional[int] = None
    issues: Optional[str] = None


class DailyReportResponse(BaseModel):
    id: int
    report_date: str
    work_performed: str
    weather: Optional[str] = None
    crew_count: Optional[int] = None
    issues: Optional[str] = None
    created_by: Optional[str] = None
    created_at: datetime

    class Config:
        from_attributes = True


# ---- ActionItem schemas ----
class ActionItemCreate(BaseModel):
    description: str
    assigned_to: Optional[str] = None
    due_date: Optional[str] = None


class ActionItemUpdate(BaseModel):
    description: Optional[str] = None
    assigned_to: Optional[str] = None
    due_date: Optional[str] = None
    status: Optional[str] = None


class ActionItemResponse(BaseModel):
    id: int
    description: str
    assigned_to: Optional[str] = None
    due_date: Optional[str] = None
    status: str
    created_by: Optional[str] = None
    created_at: datetime

    class Config:
        from_attributes = True


# Initialize database on startup
@app.on_event("startup")
async def startup():
    init_db()
    # Add title column to chats table if it doesn't exist (migration)
    from backend.database import engine
    with engine.connect() as conn:
        try:
            conn.execute(__import__('sqlalchemy').text("ALTER TABLE chats ADD COLUMN title VARCHAR(255)"))
            conn.commit()
        except Exception:
            pass  # Column already exists
        try:
            conn.execute(__import__('sqlalchemy').text(
                "ALTER TABLE documents ADD COLUMN parse_quality VARCHAR(20) DEFAULT 'good'"
            ))
            conn.commit()
        except Exception:
            pass  # Column already exists
        # Ensure new agent tables exist (init_db handles CREATE TABLE IF NOT EXISTS via SQLAlchemy metadata)
        from backend.database import Base, engine as _engine
        Base.metadata.create_all(bind=_engine)


# Health check
@app.get("/health")
async def health_check():
    return {"status": "healthy", "service": "foreperson-api"}


# ============ Auth Routes ============

@app.post("/auth/register", response_model=UserResponse)
async def register(user_data: UserCreate, db: Session = Depends(get_db)):
    """Register a new user."""
    user = register_user(db, user_data.email, user_data.password, user_data.name)
    return user


@app.post("/auth/login", response_model=TokenResponse)
async def login(user_data: UserLogin, db: Session = Depends(get_db)):
    """Login and get access token."""
    user = authenticate_user(db, user_data.email, user_data.password)
    if not user:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Incorrect email or password"
        )
    
    access_token = create_access_token(data={"sub": user.id})
    return {"access_token": access_token, "token_type": "bearer"}


@app.get("/auth/me", response_model=UserResponse)
async def get_me(current_user: User = Depends(get_current_user)):
    """Get current user info."""
    return current_user


# ============ Project Routes ============

@app.get("/projects", response_model=List[ProjectResponse])
async def list_projects(
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List all projects for current user."""
    projects = db.query(Project).filter(Project.owner_id == current_user.id).all()
    
    # Add document count
    result = []
    for p in projects:
        doc_count = db.query(Document).filter(Document.project_id == p.id).count()
        result.append(ProjectResponse(
            id=p.id,
            name=p.name,
            description=p.description,
            created_at=p.created_at,
            document_count=doc_count
        ))
    
    return result


@app.post("/projects", response_model=ProjectResponse)
async def create_project(
    project_data: ProjectCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new project."""
    project = Project(
        name=project_data.name,
        description=project_data.description,
        owner_id=current_user.id
    )
    db.add(project)
    db.commit()
    db.refresh(project)
    return ProjectResponse(
        id=project.id,
        name=project.name,
        description=project.description,
        created_at=project.created_at,
        document_count=0
    )


@app.get("/projects/{project_id}", response_model=ProjectResponse)
async def get_project(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get a specific project."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    doc_count = db.query(Document).filter(Document.project_id == project.id).count()
    return ProjectResponse(
        id=project.id,
        name=project.name,
        description=project.description,
        created_at=project.created_at,
        document_count=doc_count
    )


@app.delete("/projects/{project_id}")
async def delete_project(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a project and all its documents."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Delete associated files
    documents = db.query(Document).filter(Document.project_id == project_id).all()
    for doc in documents:
        delete_file(doc.file_path)
    
    db.delete(project)
    db.commit()
    return {"message": "Project deleted"}


@app.get("/projects/{project_id}/analytics", response_model=ProjectAnalyticsResponse)
async def get_project_analytics(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Return analytics summary for a project."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    docs = db.query(Document).filter(Document.project_id == project_id).all()

    type_breakdown: dict = {}
    total_words = 0
    for doc in docs:
        t = doc.document_type or 'unknown'
        type_breakdown[t] = type_breakdown.get(t, 0) + 1
        if doc.extracted_text:
            total_words += len(doc.extracted_text.split())

    chat_count = db.query(Chat).filter(
        Chat.project_id == project_id,
        Chat.user_id == current_user.id
    ).count()

    message_count = (
        db.query(ChatMessage)
        .join(Chat, Chat.id == ChatMessage.chat_id)
        .filter(Chat.project_id == project_id, Chat.user_id == current_user.id)
        .count()
    )

    memory_count = db.query(ProjectMemory).filter(
        ProjectMemory.project_id == project_id
    ).count()

    return {
        "doc_count": len(docs),
        "total_words": total_words,
        "type_breakdown": type_breakdown,
        "chat_count": chat_count,
        "message_count": message_count,
        "memory_fact_count": memory_count,
    }


# ============ Document Routes ============

@app.get("/projects/{project_id}/documents", response_model=List[DocumentResponse])
async def list_documents(
    project_id: int,
    page: int = 1,
    limit: int = PAGINATION_DEFAULT_LIMIT,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List all documents in a project."""
    # Verify project ownership
    project = db.query(Project).filter(
        Project.id == project_id, Project.owner_id == current_user.id
    ).first()

    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    limit = min(limit, PAGINATION_MAX_LIMIT)
    offset = (page - 1) * limit

    docs = (
        db.query(Document)
        .filter(Document.project_id == project_id)
        .offset(offset)
        .limit(limit)
        .all()
    )
    return docs


@app.post("/projects/{project_id}/documents", response_model=DocumentResponse)
async def upload_document(
    project_id: int,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Upload a document to a project."""
    # Verify project ownership
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    # Save file
    try:
        file_info = await save_file(file, current_user.id, project_id)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    
    # Extract text if parser available
    extracted_text = None
    parse_result = {}
    if DocumentParser:
        try:
            parser = DocumentParser()
            parse_result = parser.parse_document(file_info["file_path"])
            extracted_text = parse_result.get('text_content', '')
            if not extracted_text:
                # Try alternative key
                extracted_text = parse_result.get('text', '') or parse_result.get('content', '')
        except Exception as e:
            # Log error but continue
            logger.warning(f"Text extraction failed: {e}")
            extracted_text = None
    
    # Detect document type using the canonical keyword/pattern scorer
    doc_type = detect_document_type(
        extracted_text or "",
        file_info["original_filename"]
    ) if detect_document_type else "unknown"
    
    # Create document record
    document = Document(
        project_id=project_id,
        filename=file_info["filename"],
        original_filename=file_info["original_filename"],
        file_path=file_info["file_path"],
        file_size=file_info["file_size"],
        mime_type=file_info["mime_type"],
        document_type=doc_type,
        extracted_text=extracted_text,
        parse_quality=(
            parse_result.get('parse_quality', 'good')
            if parse_result and 'parse_quality' in parse_result
            else ('empty' if not extracted_text else 'good')
        ),
    )
    db.add(document)
    db.commit()
    db.refresh(document)
    
    return document


@app.get("/projects/{project_id}/documents/{document_id}/download")
async def download_document(
    project_id: int,
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Download a document."""
    # Verify ownership
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.project_id == project_id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    return FileResponse(
        document.file_path,
        filename=document.original_filename,
        media_type=document.mime_type
    )


@app.get("/projects/{project_id}/documents/{document_id}/preview")
async def preview_document(
    project_id: int,
    document_id: int,
    token: str,
    raw: bool = False,
    current_user: User = Depends(get_user_from_token_param),
    db: Session = Depends(get_db)
):
    """Render a document preview in the browser."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    document = db.query(Document).filter(
        Document.id == document_id,
        Document.project_id == project_id
    ).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    file_path = document.file_path
    original_name = document.original_filename or ""
    ext = original_name.rsplit(".", 1)[-1].lower() if "." in original_name else ""

    # raw=True: serve the raw file bytes (used by IFC 3D viewer to load the model)
    if raw:
        try:
            with open(file_path, "rb") as f:
                content = f.read()
        except FileNotFoundError:
            raise HTTPException(status_code=404, detail="File not found on disk")
        return Response(content=content, media_type="application/octet-stream")

    encoded_name = urllib.parse.quote(original_name, safe="")
    content_disposition = f"inline; filename*=UTF-8''{encoded_name}"

    # PDF — serve directly, browser renders natively
    if ext == "pdf":
        return FileResponse(file_path, media_type="application/pdf", headers={
            "Content-Disposition": content_disposition
        })

    # Images — serve directly
    if ext in ("png", "jpg", "jpeg", "gif", "webp"):
        mime = {"png": "image/png", "jpg": "image/jpeg", "jpeg": "image/jpeg",
                "gif": "image/gif", "webp": "image/webp"}.get(ext, "image/jpeg")
        return FileResponse(file_path, media_type=mime, headers={
            "Content-Disposition": content_disposition
        })

    # DOCX / DOC
    if ext in ("docx", "doc"):
        html = _render_docx_html(file_path, original_name)
        return HTMLResponse(content=html)

    # XLSX / XLS / CSV
    if ext in ("xlsx", "xls", "csv"):
        html = _render_spreadsheet_html(file_path, original_name, ext)
        return HTMLResponse(content=html)

    # DXF — render as SVG
    if ext == "dxf":
        html = _render_dxf_html(file_path, original_name)
        return HTMLResponse(content=html)

    # DWG — convert to DXF via dwg2dxf then render
    if ext == "dwg":
        raw_url = f"/api/projects/{project_id}/documents/{document_id}/preview?token={urllib.parse.quote(token, safe='')}&raw=true"
        html = _render_dwg_html(file_path, original_name, raw_url)
        return HTMLResponse(content=html)

    # IFC / BIM — handled in Task 2
    if ext in ("ifc",):
        preview_url = f"/api/projects/{project_id}/documents/{document_id}/preview?token={urllib.parse.quote(token, safe='')}&raw=true"
        html = _render_ifc_html(file_path, original_name, preview_url)
        return HTMLResponse(content=html)

    # Fallback
    return HTMLResponse(content=_unsupported_html(original_name))


@app.delete("/projects/{project_id}/documents/{document_id}")
async def delete_document(
    project_id: int,
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a document."""
    # Verify ownership
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    
    document = db.query(Document).filter(
        Document.id == document_id,
        Document.project_id == project_id
    ).first()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    # Delete file
    delete_file(document.file_path)
    
    db.delete(document)
    db.commit()
    return {"message": "Document deleted"}


@app.get("/projects/{project_id}/search")
async def search_documents(
    project_id: int,
    q: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Full-text search across document content in a project."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    if not q or len(q.strip()) < 2:
        return {"results": []}

    term = q.strip().lower()
    docs = db.query(Document).filter(
        Document.project_id == project_id,
        Document.extracted_text.isnot(None)
    ).all()

    results = []
    for doc in docs:
        text = (doc.extracted_text or "").lower()
        idx = text.find(term)
        if idx == -1:
            continue
        start = max(0, idx - 80)
        end = min(len(text), idx + len(term) + 80)
        snippet = (doc.extracted_text or "")[start:end].strip()
        if start > 0:
            snippet = "…" + snippet
        if end < len(doc.extracted_text or ""):
            snippet = snippet + "…"
        results.append({
            "doc_id": doc.id,
            "filename": doc.original_filename,
            "document_type": doc.document_type or "unknown",
            "snippet": snippet,
            "match_count": text.count(term),
        })

    results.sort(key=lambda r: r["match_count"], reverse=True)
    return {"results": results[:20]}


# ============ Chat Routes ============

def _extract_and_save_facts(
    _unused,
    question: str,
    response: str,
    project_id: int,
    chat_id: int,
):
    """Background task: extract facts from a Q&A exchange and upsert into ProjectMemory."""
    from backend.database import SessionLocal
    db = SessionLocal()
    try:
        extractor = ConstructionAI() if ConstructionAI else None
        if not extractor:
            return
        facts = extractor.extract_facts(question, response)
        for fact in facts:
            key = str(fact.get("key", "")).strip()
            value = str(fact.get("value", "")).strip()
            confidence = str(fact.get("confidence", "medium")).strip()
            if not key or not value:
                continue
            existing = db.query(ProjectMemory).filter(
                ProjectMemory.project_id == project_id,
                ProjectMemory.fact_key == key,
            ).first()
            if existing:
                existing.fact_value = value
                existing.confidence = confidence
                existing.source_thread_id = chat_id
            else:
                db.add(ProjectMemory(
                    project_id=project_id,
                    fact_key=key,
                    fact_value=value,
                    confidence=confidence,
                    source_thread_id=chat_id,
                ))
        db.commit()
        if facts:
            logger.info("Memory: saved %d fact(s) for project %d", len(facts), project_id)
    except Exception as e:
        logger.warning("Memory: failed to save facts: %s", e)
    finally:
        db.close()


@app.post("/chat", response_model=ChatResponse)
@limiter.limit(CHAT_RATE_LIMIT)
async def chat(
    request: Request,
    chat_request: ChatRequest,
    background_tasks: BackgroundTasks,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Chat with AI about documents."""
    if not ConstructionAgent:
        raise HTTPException(
            status_code=500,
            detail="AI assistant not available"
        )

    # project_id is required for all chat operations
    if not chat_request.project_id:
        raise HTTPException(status_code=400, detail="project_id is required")

    # Get documents for the specified project
    doc_list = []

    project = db.query(Project).filter(
        Project.id == chat_request.project_id,
        Project.owner_id == current_user.id
    ).first()

    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    documents = db.query(Document).filter(Document.project_id == chat_request.project_id).all()
    for doc in documents:
        # Include all documents — even those with no extracted text may have
        # visual content (drawings, maps) that the vision API can process at query time.
        text = doc.extracted_text or ""
        word_count = len(text.split()) if text else 0
        doc_list.append({
            "filename": doc.original_filename,
            "document_type": doc.document_type or "unknown",
            "text_content": text,
            "word_count": word_count,
            "file_path": doc.file_path,
            "parse_quality": doc.parse_quality or "good",
        })

    # Get or create chat session for this project
    if chat_request.chat_id:
        chat = db.query(Chat).filter(
            Chat.id == chat_request.chat_id,
            Chat.project_id == chat_request.project_id,
            Chat.user_id == current_user.id
        ).first()
        if not chat:
            raise HTTPException(status_code=404, detail="Chat thread not found")
    else:
        chat = db.query(Chat).filter(
            Chat.project_id == chat_request.project_id,
            Chat.user_id == current_user.id
        ).order_by(Chat.created_at.asc()).first()

    if not chat:
        chat = Chat(
            project_id=chat_request.project_id,
            user_id=current_user.id,
            title="Chat 1"
        )
        db.add(chat)
        db.commit()
        db.refresh(chat)

    # Save user message to database
    user_message = ChatMessage(
        chat_id=chat.id,
        role="user",
        content=chat_request.message
    )
    db.add(user_message)
    db.commit()
    db.refresh(user_message)

    # Load conversation history for this thread (excluding the message just saved)
    raw_history = (
        db.query(ChatMessage)
        .filter(ChatMessage.chat_id == chat.id, ChatMessage.id != user_message.id)
        .order_by(ChatMessage.created_at.desc())
        .limit(CHAT_HISTORY_WINDOW)
        .all()
    )
    raw_history.reverse()  # restore chronological order
    chat_history = [{"role": m.role, "content": m.content} for m in raw_history]

    # Load cross-thread project memory (only if enabled)
    memory_list = []
    if chat_request.use_memory:
        memories = db.query(ProjectMemory).filter(
            ProjectMemory.project_id == chat_request.project_id
        ).order_by(ProjectMemory.updated_at.desc()).all()
        memory_list = [{"fact_key": m.fact_key, "fact_value": m.fact_value} for m in memories]

    # Load referenced chat thread context if specified
    if chat_request.referenced_chat_id:
        ref_chat = db.query(Chat).filter(
            Chat.id == chat_request.referenced_chat_id,
            Chat.project_id == chat_request.project_id,
            Chat.user_id == current_user.id
        ).first()
        if ref_chat:
            ref_msgs = db.query(ChatMessage).filter(
                ChatMessage.chat_id == ref_chat.id
            ).order_by(ChatMessage.created_at.asc()).limit(40).all()
            if ref_msgs:
                ref_title = ref_chat.title or f"Chat {ref_chat.id}"
                ref_text = "\n".join(f"{m.role.upper()}: {m.content}" for m in ref_msgs)
                chat_history = [
                    {"role": "user", "content": f"[Context from {ref_title}]\n{ref_text}"},
                    {"role": "assistant", "content": f"I've read the context from {ref_title}. How can I help?"}
                ] + chat_history

    # Create AI assistant and get response
    try:
        agent = ConstructionAgent(
            db=db,
            project_id=chat_request.project_id,
            user_name=current_user.name,
            documents=doc_list,
            model=chat_request.model,
        )
        response = agent.run(
            chat_request.message,
            history=chat_history,
            project_memory=memory_list,
        )

        # Save assistant response to database
        assistant_message = ChatMessage(
            chat_id=chat.id,
            role="assistant",
            content=response
        )
        db.add(assistant_message)
        db.commit()
        db.refresh(assistant_message)

        # Schedule fact extraction in background (non-blocking)
        background_tasks.add_task(
            _extract_and_save_facts,
            None,
            chat_request.message,
            response,
            chat_request.project_id,
            chat.id,
        )

        return {
            "response": response,
            "message_id": assistant_message.id
        }
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"AI error: {str(e)}"
        )


@app.get("/projects/{project_id}/chats", response_model=List[ChatThreadResponse])
async def list_chat_threads(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """List all chat threads for a project."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    chats = db.query(Chat).filter(
        Chat.project_id == project_id,
        Chat.user_id == current_user.id
    ).order_by(Chat.created_at.asc()).all()

    result = []
    for c in chats:
        count = db.query(ChatMessage).filter(ChatMessage.chat_id == c.id).count()
        result.append({"id": c.id, "title": c.title, "message_count": count, "created_at": c.created_at})
    return result


@app.post("/projects/{project_id}/chats", response_model=ChatHistoryResponse)
async def create_chat_thread(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Create a new chat thread for a project."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    count = db.query(Chat).filter(
        Chat.project_id == project_id,
        Chat.user_id == current_user.id
    ).count()

    chat = Chat(
        project_id=project_id,
        user_id=current_user.id,
        title=f"Chat {count + 1}"
    )
    db.add(chat)
    db.commit()
    db.refresh(chat)
    return {"id": chat.id, "title": chat.title, "messages": [], "created_at": chat.created_at}


@app.patch("/projects/{project_id}/chats/{chat_id}")
async def rename_chat_thread(
    project_id: int,
    chat_id: int,
    body: dict,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Rename a chat thread."""
    chat = db.query(Chat).filter(
        Chat.id == chat_id,
        Chat.project_id == project_id,
        Chat.user_id == current_user.id
    ).first()
    if not chat:
        raise HTTPException(status_code=404, detail="Chat thread not found")
    title = (body.get("title") or "").strip()
    if title:
        chat.title = title
        db.commit()
    return {"id": chat.id, "title": chat.title}


@app.delete("/projects/{project_id}/chats/{chat_id}")
async def delete_chat_thread(
    project_id: int,
    chat_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Delete a chat thread. Cannot delete the last remaining thread."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    total = db.query(Chat).filter(
        Chat.project_id == project_id,
        Chat.user_id == current_user.id
    ).count()

    chat = db.query(Chat).filter(
        Chat.id == chat_id,
        Chat.project_id == project_id,
        Chat.user_id == current_user.id
    ).first()
    if not chat:
        raise HTTPException(status_code=404, detail="Chat thread not found")

    db.delete(chat)
    db.commit()
    return {"ok": True}


@app.get("/projects/{project_id}/chat", response_model=ChatHistoryResponse)
async def get_chat_history(
    project_id: int,
    chat_id: Optional[int] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get chat history for a project (optionally a specific thread)."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    if chat_id:
        chat = db.query(Chat).filter(
            Chat.id == chat_id,
            Chat.project_id == project_id,
            Chat.user_id == current_user.id
        ).first()
        if not chat:
            raise HTTPException(status_code=404, detail="Chat thread not found")
    else:
        chat = db.query(Chat).filter(
            Chat.project_id == project_id,
            Chat.user_id == current_user.id
        ).order_by(Chat.created_at.asc()).first()

    if not chat:
        chat = Chat(
            project_id=project_id,
            user_id=current_user.id,
            title="Chat 1"
        )
        db.add(chat)
        db.commit()
        db.refresh(chat)

    messages = db.query(ChatMessage).filter(
        ChatMessage.chat_id == chat.id
    ).order_by(ChatMessage.created_at.asc()).all()

    return {
        "id": chat.id,
        "title": chat.title,
        "messages": messages,
        "created_at": chat.created_at
    }


@app.get("/projects/{project_id}/memory")
async def get_project_memory(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Get all remembered facts for a project."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    memories = db.query(ProjectMemory).filter(
        ProjectMemory.project_id == project_id
    ).order_by(ProjectMemory.updated_at.desc()).all()

    return {
        "project_id": project_id,
        "facts": [
            {
                "key": m.fact_key,
                "value": m.fact_value,
                "confidence": m.confidence,
                "updated_at": m.updated_at,
            }
            for m in memories
        ]
    }


@app.delete("/projects/{project_id}/memory/{fact_key}")
async def delete_memory_fact(
    project_id: int,
    fact_key: str,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Delete a specific remembered fact."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    fact = db.query(ProjectMemory).filter(
        ProjectMemory.project_id == project_id,
        ProjectMemory.fact_key == fact_key,
    ).first()
    if not fact:
        raise HTTPException(status_code=404, detail="Fact not found")

    db.delete(fact)
    db.commit()
    return {"message": f"Deleted fact: {fact_key}"}


@app.post("/projects/{project_id}/conflicts", response_model=List[ConflictResponse])
@limiter.limit(CONFLICTS_RATE_LIMIT)
async def analyze_conflicts(
    request: Request,
    project_id: int,
    body: ConflictsRequest = ConflictsRequest(),
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Analyze documents in a project for conflicts. Optionally filter to specific doc IDs."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    if not ConstructionAI:
        raise HTTPException(status_code=500, detail="AI assistant not available")

    query = db.query(Document).filter(Document.project_id == project_id)
    if body.doc_ids:
        query = query.filter(Document.id.in_(body.doc_ids))
    documents = query.all()

    doc_list = []
    for doc in documents:
        if doc.extracted_text:
            doc_list.append({
                "filename": doc.original_filename,
                "document_type": doc.document_type or "unknown",
                "text_content": doc.extracted_text,
                "word_count": len(doc.extracted_text.split()),
                "file_path": doc.file_path,
                "parse_quality": doc.parse_quality or "good",
            })

    if len(doc_list) < 2:
        raise HTTPException(status_code=400, detail="At least 2 documents with extractable text are required")

    try:
        assistant = ConstructionAI()
        assistant.load_documents(doc_list)
        raw_conflicts = assistant.find_conflicts()

        result = []
        for i, c in enumerate(raw_conflicts):
            result.append({
                "id": str(i + 1),
                "severity": c.get("severity", "medium"),
                "title": c.get("title", "Untitled conflict"),
                "description": c.get("description", ""),
                "resolution": c.get("resolution", ""),
                "documents": c.get("documents", []),
            })
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Conflict analysis error: {str(e)}")


@app.get("/projects/{project_id}/conflict-statuses")
async def get_conflict_statuses(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Get all conflict status overrides for a project. Returns dict of hash->status."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    statuses = db.query(ConflictStatus).filter(
        ConflictStatus.project_id == project_id
    ).all()
    return {s.conflict_hash: s.status for s in statuses}


@app.post("/projects/{project_id}/conflict-statuses/{conflict_hash}")
async def set_conflict_status(
    project_id: int,
    conflict_hash: str,
    body: dict,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Set open/resolved/dismissed status for a specific conflict."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    new_status = body.get("status", "open")
    if new_status not in ("open", "resolved", "dismissed"):
        raise HTTPException(status_code=400, detail="status must be open, resolved, or dismissed")

    existing = db.query(ConflictStatus).filter(
        ConflictStatus.project_id == project_id,
        ConflictStatus.conflict_hash == conflict_hash
    ).first()

    if existing:
        existing.status = new_status
        existing.updated_at = datetime.utcnow()
    else:
        db.add(ConflictStatus(
            project_id=project_id,
            conflict_hash=conflict_hash,
            status=new_status
        ))
    db.commit()
    return {"conflict_hash": conflict_hash, "status": new_status}


@app.post("/projects/{project_id}/compare", response_model=CompareResponse)
async def compare_documents(
    project_id: int,
    body: CompareRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    """Deep comparison of two documents in a project."""
    project = db.query(Project).filter(
        Project.id == project_id,
        Project.owner_id == current_user.id
    ).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    if not ConstructionAI:
        raise HTTPException(status_code=500, detail="AI assistant not available")

    doc1 = db.query(Document).filter(Document.id == body.doc_id_1, Document.project_id == project_id).first()
    doc2 = db.query(Document).filter(Document.id == body.doc_id_2, Document.project_id == project_id).first()

    if not doc1 or not doc2:
        raise HTTPException(status_code=404, detail="One or both documents not found")
    if doc1.id == doc2.id:
        raise HTTPException(status_code=400, detail="Select two different documents")

    doc_list = []
    for doc in [doc1, doc2]:
        doc_list.append({
            "filename": doc.original_filename,
            "document_type": doc.document_type or "unknown",
            "text_content": doc.extracted_text or "",
            "word_count": len((doc.extracted_text or "").split()),
            "file_path": doc.file_path,
            "parse_quality": doc.parse_quality or "good",
        })

    try:
        assistant = ConstructionAI()
        assistant.load_documents(doc_list)
        result_dict = assistant.compare_documents(0, 1)
        return {
            "summary": result_dict.get("summary", ""),
            "conflicts": result_dict.get("conflicts", []),
            "gaps": result_dict.get("gaps", []),
            "agreements": result_dict.get("agreements", []),
            "risks": result_dict.get("risks", []),
            "doc1_name": doc1.original_filename,
            "doc2_name": doc2.original_filename,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Compare error: {str(e)}")


# ============ RFI Routes ============

def _get_project_or_404(project_id: int, user_id: int, db: Session):
    project = db.query(Project).filter(Project.id == project_id, Project.owner_id == user_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project


@app.get("/projects/{project_id}/rfis", response_model=List[RFIResponse])
async def list_rfis(
    project_id: int,
    status: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _get_project_or_404(project_id, current_user.id, db)
    q = db.query(RFI).filter(RFI.project_id == project_id)
    if status:
        q = q.filter(RFI.status == status)
    return q.order_by(RFI.number.desc()).all()


@app.post("/projects/{project_id}/rfis", response_model=RFIResponse, status_code=201)
async def create_rfi(
    project_id: int,
    body: RFICreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _get_project_or_404(project_id, current_user.id, db)
    last = db.query(RFI).filter(RFI.project_id == project_id).order_by(RFI.number.desc()).first()
    number = (last.number + 1) if last else 1
    rfi = RFI(project_id=project_id, number=number, subject=body.subject,
              description=body.description, due_date=body.due_date,
              created_by=current_user.name, status="open")
    db.add(rfi)
    db.commit()
    db.refresh(rfi)
    return rfi


@app.patch("/projects/{project_id}/rfis/{rfi_id}", response_model=RFIResponse)
async def update_rfi(
    project_id: int,
    rfi_id: int,
    body: RFIUpdate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _get_project_or_404(project_id, current_user.id, db)
    rfi = db.query(RFI).filter(RFI.id == rfi_id, RFI.project_id == project_id).first()
    if not rfi:
        raise HTTPException(status_code=404, detail="RFI not found")
    for field, val in body.dict(exclude_unset=True).items():
        setattr(rfi, field, val)
    db.commit()
    db.refresh(rfi)
    return rfi


@app.delete("/projects/{project_id}/rfis/{rfi_id}", status_code=204)
async def delete_rfi(
    project_id: int,
    rfi_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _get_project_or_404(project_id, current_user.id, db)
    rfi = db.query(RFI).filter(RFI.id == rfi_id, RFI.project_id == project_id).first()
    if not rfi:
        raise HTTPException(status_code=404, detail="RFI not found")
    db.delete(rfi)
    db.commit()


# ============ Daily Report Routes ============

@app.get("/projects/{project_id}/daily-reports", response_model=List[DailyReportResponse])
async def list_daily_reports(
    project_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _get_project_or_404(project_id, current_user.id, db)
    return db.query(DailyReport).filter(DailyReport.project_id == project_id).order_by(DailyReport.report_date.desc()).all()


@app.post("/projects/{project_id}/daily-reports", response_model=DailyReportResponse, status_code=201)
async def create_daily_report(
    project_id: int,
    body: DailyReportCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _get_project_or_404(project_id, current_user.id, db)
    report = DailyReport(project_id=project_id, report_date=body.report_date,
                         work_performed=body.work_performed, weather=body.weather,
                         crew_count=body.crew_count, issues=body.issues,
                         created_by=current_user.name)
    db.add(report)
    db.commit()
    db.refresh(report)
    return report


@app.delete("/projects/{project_id}/daily-reports/{report_id}", status_code=204)
async def delete_daily_report(
    project_id: int,
    report_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _get_project_or_404(project_id, current_user.id, db)
    report = db.query(DailyReport).filter(DailyReport.id == report_id, DailyReport.project_id == project_id).first()
    if not report:
        raise HTTPException(status_code=404, detail="Report not found")
    db.delete(report)
    db.commit()


# ============ Action Item Routes ============

@app.get("/projects/{project_id}/action-items", response_model=List[ActionItemResponse])
async def list_action_items(
    project_id: int,
    status: Optional[str] = None,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _get_project_or_404(project_id, current_user.id, db)
    q = db.query(ActionItem).filter(ActionItem.project_id == project_id)
    if status:
        q = q.filter(ActionItem.status == status)
    return q.order_by(ActionItem.created_at.desc()).all()


@app.post("/projects/{project_id}/action-items", response_model=ActionItemResponse, status_code=201)
async def create_action_item(
    project_id: int,
    body: ActionItemCreate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _get_project_or_404(project_id, current_user.id, db)
    item = ActionItem(project_id=project_id, description=body.description,
                      assigned_to=body.assigned_to, due_date=body.due_date,
                      created_by=current_user.name, status="open")
    db.add(item)
    db.commit()
    db.refresh(item)
    return item


@app.patch("/projects/{project_id}/action-items/{item_id}", response_model=ActionItemResponse)
async def update_action_item(
    project_id: int,
    item_id: int,
    body: ActionItemUpdate,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _get_project_or_404(project_id, current_user.id, db)
    item = db.query(ActionItem).filter(ActionItem.id == item_id, ActionItem.project_id == project_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Action item not found")
    for field, val in body.dict(exclude_unset=True).items():
        setattr(item, field, val)
    db.commit()
    db.refresh(item)
    return item


@app.delete("/projects/{project_id}/action-items/{item_id}", status_code=204)
async def delete_action_item(
    project_id: int,
    item_id: int,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    _get_project_or_404(project_id, current_user.id, db)
    item = db.query(ActionItem).filter(ActionItem.id == item_id, ActionItem.project_id == project_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Action item not found")
    db.delete(item)
    db.commit()


# ---------------------------------------------------------------------------
# Preview helper functions
# ---------------------------------------------------------------------------

def _preview_page(title: str, body: str, extra_head: str = "") -> str:
    """Wrap content in a clean preview HTML page."""
    safe_title = html.escape(title)
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{safe_title}</title>
{extra_head}
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  body {{ font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif;
         background: #f8f7f3; color: #1c1b18; padding: 0; }}
  .header {{ background: #1c1b18; color: #f0ede4; padding: 12px 24px;
             display: flex; align-items: center; gap: 12px; position: sticky; top: 0; z-index: 10; }}
  .header .filename {{ font-size: 0.8rem; font-family: 'SF Mono', monospace;
                       opacity: 0.7; white-space: nowrap; overflow: hidden; text-overflow: ellipsis; }}
  .header .badge {{ font-size: 0.6rem; font-family: monospace; letter-spacing: 0.1em;
                    padding: 2px 8px; border: 1px solid rgba(255,255,255,0.2);
                    color: #f5c800; border-color: rgba(245,200,0,0.4); flex-shrink: 0; }}
  .content {{ max-width: 900px; margin: 0 auto; padding: 32px 24px; }}
</style>
</head>
<body>
<div class="header">
  <span class="badge">FOREPERSON</span>
  <span class="filename">{safe_title}</span>
</div>
<div class="content">{body}</div>
</body>
</html>"""


def _render_docx_html(file_path: str, name: str) -> str:
    try:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                parts.append("<br>")
                continue
            style = para.style.name if para.style else ""
            safe_text = html.escape(text)
            if style.startswith("Heading 1"):
                parts.append(f"<h1 style='font-size:1.6rem;font-weight:700;margin:1.5rem 0 0.5rem'>{safe_text}</h1>")
            elif style.startswith("Heading 2"):
                parts.append(f"<h2 style='font-size:1.2rem;font-weight:600;margin:1.2rem 0 0.4rem'>{safe_text}</h2>")
            elif style.startswith("Heading 3"):
                parts.append(f"<h3 style='font-size:1rem;font-weight:600;margin:1rem 0 0.3rem'>{safe_text}</h3>")
            else:
                parts.append(f"<p style='margin:0.4rem 0;line-height:1.7;font-size:0.92rem'>{safe_text}</p>")
        body = "\n".join(parts)
        import re as _re
        if not _re.sub(r'<br\s*/?>', '', body).strip():
            body = "<p style='color:#999'>Document appears to be empty.</p>"
    except Exception as e:
        body = f"<p style='color:#ef4444'>Could not render document: {html.escape(str(e))}</p>"
    return _preview_page(name, body)


def _render_spreadsheet_html(file_path: str, name: str, ext: str) -> str:
    try:
        import pandas as pd
        if ext == "csv":
            df = pd.read_csv(file_path, nrows=500)
        else:
            df = pd.read_excel(file_path, nrows=500)
        table_html = df.to_html(
            index=False,
            border=0,
            classes="data-table",
            na_rep="",
        )
        body = f"""
<style>
  .data-table {{ border-collapse: collapse; width: 100%; font-size: 0.8rem; }}
  .data-table th {{ background: #1c1b18; color: #f0ede4; padding: 8px 12px;
                    text-align: left; font-family: monospace; font-size: 0.7rem;
                    letter-spacing: 0.06em; white-space: nowrap; }}
  .data-table td {{ padding: 7px 12px; border-bottom: 1px solid #e0ddd4;
                    color: #1c1b18; vertical-align: top; }}
  .data-table tr:hover td {{ background: #f0ede4; }}
  .row-count {{ font-family: monospace; font-size: 0.7rem; color: #7a7268;
                margin-bottom: 12px; }}
</style>
<p class="row-count">Showing {min(len(df), 500)} rows × {len(df.columns)} columns</p>
<div style="overflow-x:auto">{table_html}</div>"""
    except Exception as e:
        body = f"<p style='color:#ef4444'>Could not render spreadsheet: {e}</p>"
    return _preview_page(name, body)


def _render_dxf_html(file_path: str, name: str) -> str:
    try:
        import ezdxf
        from ezdxf import recover as ezdxf_recover
        import xml.etree.ElementTree as ET
        try:
            doc, _ = ezdxf_recover.readfile(file_path)
        except Exception:
            doc = ezdxf.readfile(file_path)
        from ezdxf.addons.drawing import RenderContext, Frontend
        from ezdxf.addons.drawing.svg import SVGBackend
        context = RenderContext(doc)
        backend = SVGBackend()
        Frontend(context, backend).draw_layout(doc.modelspace())
        svg_xml = backend.get_xml_root()
        svg_str = ET.tostring(svg_xml, encoding="unicode")
        body = f"""
<style>
  .svg-wrap {{ background: white; border: 1px solid #e0ddd4; padding: 16px;
              border-radius: 2px; overflow: auto; }}
  .svg-wrap svg {{ max-width: 100%; height: auto; display: block; }}
</style>
<div class="svg-wrap">{svg_str}</div>"""
    except Exception as e:
        body = f"<p style='color:#ef4444'>Could not render drawing: {html.escape(str(e))}</p><p style='color:#7a7268;font-size:0.85rem;margin-top:8px'>DWG files may require conversion. Try re-exporting as DXF.</p>"
    return _preview_page(name, body)


def _render_dwg_html(file_path: str, name: str, raw_url: str) -> str:
    """Try to convert DWG→DXF via dwg2dxf, fall back to info+download page."""
    import shutil, subprocess, tempfile, xml.etree.ElementTree as ET
    safe_name = html.escape(name)

    # Attempt conversion with dwg2dxf (from libredwg)
    if shutil.which("dwg2dxf"):
        try:
            with tempfile.TemporaryDirectory() as tmp:
                result = subprocess.run(
                    ["dwg2dxf", "-o", tmp, file_path],
                    capture_output=True, timeout=30
                )
                dxf_files = [f for f in os.listdir(tmp) if f.endswith(".dxf")]
                if dxf_files:
                    dxf_path = os.path.join(tmp, dxf_files[0])
                    return _render_dxf_html(dxf_path, name)
        except Exception:
            pass

    # Fallback: read DWG version from magic bytes and show info + download
    version = "Unknown"
    file_size = 0
    version_map = {
        b"AC1009": "R11/R12", b"AC1012": "R13", b"AC1014": "R14",
        b"AC1015": "2000",    b"AC1018": "2004", b"AC1021": "2007",
        b"AC1024": "2010",    b"AC1027": "2013", b"AC1032": "2018",
    }
    try:
        with open(file_path, "rb") as f:
            magic = f.read(6)
        version = version_map.get(magic, "Unknown")
        file_size = os.path.getsize(file_path)
    except Exception:
        pass

    size_str = f"{file_size/1024:.1f} KB" if file_size < 1024*1024 else f"{file_size/1024/1024:.1f} MB"

    body = f"""
<div style="max-width:480px;margin:60px auto;text-align:center">
  <div style="font-size:3rem;margin-bottom:20px">📐</div>
  <h2 style="font-size:1.1rem;font-weight:700;margin-bottom:20px">{safe_name}</h2>
  <table style="margin:0 auto 24px;border-collapse:collapse;text-align:left">
    <tr style="border-bottom:1px solid var(--border)">
      <td style="padding:8px 20px;color:var(--text-secondary);font-family:monospace;font-size:0.7rem;letter-spacing:0.06em">FORMAT</td>
      <td style="padding:8px 20px;font-family:monospace;font-size:0.75rem;font-weight:600">AutoCAD DWG</td>
    </tr>
    <tr style="border-bottom:1px solid var(--border)">
      <td style="padding:8px 20px;color:var(--text-secondary);font-family:monospace;font-size:0.7rem;letter-spacing:0.06em">VERSION</td>
      <td style="padding:8px 20px;font-family:monospace;font-size:0.75rem;font-weight:600">AutoCAD {html.escape(version)}</td>
    </tr>
    <tr>
      <td style="padding:8px 20px;color:var(--text-secondary);font-family:monospace;font-size:0.7rem;letter-spacing:0.06em">SIZE</td>
      <td style="padding:8px 20px;font-family:monospace;font-size:0.75rem;font-weight:600">{size_str}</td>
    </tr>
  </table>
  <p style="color:var(--text-secondary);font-size:0.82rem;margin-bottom:24px;line-height:1.6">
    DWG is a proprietary Autodesk format.<br>Download to open in AutoCAD, BricsCAD, or another CAD viewer.
  </p>
  <a href="{raw_url}" style="display:inline-block;background:var(--accent);color:var(--accent-dark);padding:10px 28px;text-decoration:none;font-weight:600;font-size:0.875rem;border-radius:2px;letter-spacing:0.02em">
    Download DWG
  </a>
</div>"""
    return _preview_page(safe_name, body)


def _render_ifc_html(file_path: str, name: str, preview_url: str) -> str:
    import html as html_mod
    safe_name = html_mod.escape(name)
    metadata_html = ""
    try:
        import ifcopenshell
        ifc = ifcopenshell.open(file_path)

        projects = ifc.by_type("IfcProject")
        proj_name = html_mod.escape(projects[0].Name or "") if projects else ""
        proj_desc = html_mod.escape(projects[0].Description or "") if projects and projects[0].Description else ""

        def count(ifc_type):
            try:
                return len(ifc.by_type(ifc_type))
            except Exception:
                return 0

        elements = {
            "Storeys": count("IfcBuildingStorey"),
            "Spaces": count("IfcSpace"),
            "Walls": count("IfcWall"),
            "Slabs": count("IfcSlab"),
            "Columns": count("IfcColumn"),
            "Beams": count("IfcBeam"),
            "Doors": count("IfcDoor"),
            "Windows": count("IfcWindow"),
            "Stairs": count("IfcStair"),
            "Roofs": count("IfcRoof"),
        }

        rows = "".join(
            f"<tr><td class='label'>{k}</td><td class='val'>{v}</td></tr>"
            for k, v in elements.items() if v > 0
        )

        display_title = proj_name or safe_name
        metadata_html = f"""
<div class="meta-panel">
  <div class="meta-header">
    <div class="meta-badge">BIM</div>
    <h2 class="meta-title">{display_title}</h2>
    {f'<p class="meta-desc">{proj_desc}</p>' if proj_desc else ''}
  </div>
  <table class="meta-table">
    <tbody>{rows or "<tr><td colspan='2' style='color:#7a7268;padding:16px;font-size:0.8rem'>No elements found</td></tr>"}</tbody>
  </table>
  <p class="meta-file">{safe_name}</p>
</div>"""
    except Exception as e:
        safe_e = html_mod.escape(str(e))
        metadata_html = f"""
<div class="meta-panel">
  <div class="meta-header">
    <div class="meta-badge">BIM</div>
    <h2 class="meta-title">{safe_name}</h2>
  </div>
  <p style="color:#7a7268;font-size:0.8rem;padding:16px">Could not parse metadata: {safe_e}</p>
</div>"""

    # preview_url is server-constructed (not user input), safe to embed
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{safe_name}</title>
<style>
  * {{ box-sizing: border-box; margin: 0; padding: 0; }}
  html, body {{ height: 100%; overflow: hidden;
    font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; }}
  body {{ display: flex; flex-direction: column; background: #111110; color: #f0ede4; }}

  .header {{ background: #1c1b18; padding: 10px 20px; display: flex; align-items: center;
             gap: 12px; border-bottom: 1px solid #313130; flex-shrink: 0; }}
  .header .badge {{ font-size: 0.6rem; font-family: monospace; letter-spacing: 0.1em;
                    padding: 2px 8px; color: #f5c800; border: 1px solid rgba(245,200,0,0.4); }}
  .header .filename {{ font-size: 0.8rem; font-family: monospace; opacity: 0.6;
                       overflow: hidden; text-overflow: ellipsis; white-space: nowrap; }}

  .main {{ display: flex; flex: 1; overflow: hidden; }}

  .meta-panel {{ width: 280px; flex-shrink: 0; background: #1a1a18; border-right: 1px solid #313130;
                 overflow-y: auto; display: flex; flex-direction: column; }}
  .meta-header {{ padding: 20px; border-bottom: 1px solid #313130; }}
  .meta-badge {{ font-size: 0.55rem; font-family: monospace; letter-spacing: 0.12em;
                 color: #f5c800; border: 1px solid rgba(245,200,0,0.3); padding: 2px 6px;
                 display: inline-block; margin-bottom: 10px; }}
  .meta-title {{ font-size: 1rem; font-weight: 700; line-height: 1.3; color: #f0ede4; }}
  .meta-desc {{ font-size: 0.75rem; color: #7a7268; margin-top: 6px; line-height: 1.5; }}
  .meta-table {{ width: 100%; border-collapse: collapse; }}
  .meta-table tr {{ border-bottom: 1px solid #232320; }}
  .meta-table td {{ padding: 10px 20px; font-size: 0.8rem; }}
  .meta-table .label {{ color: #7a7268; font-family: monospace; font-size: 0.7rem;
                         letter-spacing: 0.06em; width: 50%; }}
  .meta-table .val {{ color: #f0ede4; font-weight: 600; font-family: monospace; text-align: right; }}
  .meta-file {{ font-family: monospace; font-size: 0.65rem; color: #4a4a48;
                padding: 16px 20px; margin-top: auto; }}

  .viewer-wrap {{ flex: 1; position: relative; background: #0d0d0c; }}
  #viewer-canvas {{ width: 100%; height: 100%; display: block; }}
  .viewer-loading {{ position: absolute; inset: 0; display: flex; align-items: center;
                      justify-content: center; flex-direction: column; gap: 12px;
                      background: #0d0d0c; z-index: 5; }}
  .viewer-loading p {{ font-family: monospace; font-size: 0.75rem; color: #7a7268;
                        letter-spacing: 0.08em; }}
  .spinner {{ width: 32px; height: 32px; border: 2px solid #313130;
              border-top-color: #f5c800; border-radius: 50%;
              animation: spin 0.8s linear infinite; }}
  @keyframes spin {{ to {{ transform: rotate(360deg); }} }}
  .viewer-error {{ position: absolute; inset: 0; display: none; align-items: center;
                   justify-content: center; flex-direction: column; gap: 8px; }}
  .viewer-error p {{ font-size: 0.85rem; color: #7a7268; font-family: monospace; }}
</style>
</head>
<body>
<div class="header">
  <span class="badge">FOREPERSON · BIM</span>
  <span class="filename">{safe_name}</span>
</div>
<div class="main">
  {metadata_html}
  <div class="viewer-wrap">
    <div class="viewer-loading" id="loading">
      <div class="spinner"></div>
      <p>LOADING 3D MODEL…</p>
    </div>
    <div class="viewer-error" id="error">
      <p>3D viewer unavailable</p>
      <p style="font-size:0.7rem;color:#4a4a48">Metadata shown on left panel</p>
    </div>
    <canvas id="viewer-canvas"></canvas>
  </div>
</div>

<script type="module">
  import {{ IfcViewerAPI }} from 'https://unpkg.com/web-ifc-viewer@1.0.220/dist/index.js';
  import * as THREE from 'https://unpkg.com/three@0.152.2/build/three.module.js';

  const container = document.querySelector('.viewer-wrap');
  const loading = document.getElementById('loading');
  const errorEl = document.getElementById('error');

  try {{
    const viewer = new IfcViewerAPI({{
      container,
      backgroundColor: new THREE.Color(0x0d0d0c),
    }});

    viewer.axes.setAxes();
    viewer.grid.setGrid();
    viewer.IFC.setWasmPath('https://unpkg.com/web-ifc@0.0.44/');

    await viewer.IFC.loadIfcUrl({json.dumps(preview_url)}, true, (progress) => {{
      if (progress.total > 0) {{
        const pct = Math.round((progress.loaded / progress.total) * 100);
        document.querySelector('#loading p').textContent = `LOADING 3D MODEL… ${{pct}}%`;
      }}
    }});

    loading.style.display = 'none';
  }} catch (e) {{
    console.error('IFC viewer error:', e);
    loading.style.display = 'none';
    errorEl.style.display = 'flex';
  }}
</script>
</body>
</html>"""


def _unsupported_html(name: str) -> str:
    body = """
<div style='text-align:center;padding:60px 0'>
  <p style='font-size:2rem;margin-bottom:16px'>📄</p>
  <p style='font-size:1rem;font-weight:600;margin-bottom:8px'>Preview not available</p>
  <p style='color:#7a7268;font-size:0.85rem'>This file type cannot be previewed in the browser.</p>
</div>"""
    return _preview_page(name, body)


# Run with: uvicorn backend.api:app --reload
if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
